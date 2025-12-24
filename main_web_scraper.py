from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from dotenv import load_dotenv
import requests
from docx import Document
from bs4 import BeautifulSoup
import csv
import re
from datetime import datetime
import os
import asyncio
from crawl4ai import *
import urllib3
import ollama
import json

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

load_dotenv()

os.makedirs("pdfs", exist_ok=True)

def save_pdf(url, folder="pdfs"):
    try:
        r = requests.get(url, verify=False, timeout=80, stream=True)
        r.raise_for_status()
        if "application/pdf" in r.headers.get("Content-Type", "").lower():
            # Create filename from last part of URL or fallback timestamp
            name = url.split("/")[-1]
            if not name.endswith(".pdf"):
                name += ".pdf"
            filepath = os.path.join(folder, name)
            with open(filepath, "wb") as f:
                for chunk in r.iter_content(8192):
                    f.write(chunk)
            print(f"✅ Saved PDF: {filepath}")
            return True
        else:
            print(f"⚠️ Not a PDF: {url}")
            return False
    except Exception as e:
        print(f"❌ Failed to download {url}: {e}")
        return False

def extract_main_content(soup):
    # Try <main> first
    main = soup.find("main")
    if main:
        return main
    # Try common content containers
    for selector in ["main#main", "main#main-content", "div#content", "div.main-content", "div#main", "article", "div.header","div.container", "div.content", "div.field-items", "section"]:
        node = soup.select_one(selector)
        if node:
            return node
    # Fallback: whole body minus header/footer/nav
    body = soup.body
    if body:
        for tag in body.find_all(["header", "footer", "nav"]):
            tag.decompose()
        return body
    return None

def fetch_body_content(url):
    try:
        r = requests.get(url, verify=False, timeout=80)
        r.raise_for_status()

        # Direct PDF?
        if "application/pdf" in r.headers.get("Content-Type", "").lower():
            save_pdf(url)
            return {"text": "", "tables": []}

        soup = BeautifulSoup(r.text, "html.parser")
        main = extract_main_content(soup)
        if not main:
            return {"text": "", "tables": []}

        # --- Extract PDF links inside page ---
        # from urllib.parse import urljoin
        # for a in main.find_all("a", href=True):
        #     href = a["href"]
        #     if href.lower().endswith(".pdf"):
        #         absolute_url = urljoin(url, href)
        #         save_pdf(absolute_url)

        # --- Extract clean text (no duplicates) ---
        texts = []
        for tag in main.find_all(["p", "ul", "li", "h1", "h2", "h3", "h4", "h5", "h6"]):
            if tag.find_parent("table"):
                continue
            text = tag.get_text(strip=True)
            if text:
                texts.append(text)

        # --- Extract tables ---
        tables = []
        for table in main.find_all("table"):
            rows, max_cols = [], 0
            for tr in table.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
                    max_cols = max(max_cols, len(cells))
            if rows:
                tables.append({"rows": rows, "cols": max_cols})

        return {"text": "\n".join(texts), "tables": tables}

    except Exception as e:
        print(f"❌ Failed to fetch {url}: {e}")
        return {"text": "", "tables": []}
    
def save_docx(links, filename="chennai_port_output.docx"):
    doc = Document()
    for url in links:
        doc.add_heading(url, level=2)
        content = fetch_body_content(url)

        # Add text
        if content["text"]:
            doc.add_paragraph(content["text"])

        # Add tables
        for table in content["tables"]:
            rows, cols = len(table["rows"]), table["cols"]
            doc_table = doc.add_table(rows=rows, cols=cols)
            for i, row in enumerate(table["rows"]):
                for j in range(cols):
                    cell_text = row[j] if j < len(row) else ""  # pad missing cells
                    doc_table.cell(i, j).text = cell_text

    doc.save(filename)
    print(f"📄 Saved DOCX: {filename}")


async def urls():
    async with AsyncWebCrawler() as crawler:
        result = await crawler.arun(
            url="https://www.cochinport.gov.in/",
        )
        return result.markdown

# async def urls():
#     async with AsyncWebCrawler() as crawler:
#         result = await crawler.arun(
#             url="https://www.chennaiport.gov.in/home",
#         )
#         return result.markdown

async def prompter():
    # Initialize the Groq LLM
    llm = ChatGroq(
    model_name="llama-3.3-70b-versatile",
    temperature=0.7
    )
    # Define the expected JSON structure
    parser = JsonOutputParser(pydantic_object={
    "type": "object",
    "properties": {
        "name": {"type": "string"},
        "price": {"type": "number"},
        "features": {
            "type": "array",
            "items": {"type": "string"}
        }
    }
    })

    url_content = await urls()
    # Create a prompt template
    # prompt = ChatPromptTemplate.from_messages([
    # (
    #     "system",
    #     """
    #     You are an intelligent URL analyzer. Your job is to examine the provided Markdown content (from the crawled site) and return ONLY the list of URLs that lead (or clearly redirect) to pages that provide information about the website or the organization that runs it.

    #     Selection criteria (use surrounding link text and nearby content to decide):
    #     - Include links whose target pages are clearly "About", "About Us", "Who we are", "Our mission", "Team", "Contact", "Company", "Corporate", "Governance", "History", "Legal/Privacy" or any page that describes the organization or site operator.
    #     - Exclude links to blog posts, news, product pages, documentation, login, search results, external social media profiles, or purely transactional pages unless they clearly contain organization info.
    #     - If a link is relative, convert it to an absolute URL using the base domain of the input URL.
    #     - Only include URLs that point to content pages (HTTP/HTTPS), not anchors like "#section" or javascript: links.

    #     Important output rules (strictly follow):
    #     1) Output MUST be valid JSON and NOTHING else.
    #     2) The JSON must be exactly: {{"important_links": ["https://...", "https://...", ...]}}
    #     3) Use absolute, fully qualified URLs only.
    #     4) Return an empty list if no qualifying links are found.

    #     Examples:
    #     - Correct: {{"important_links": ["https://www.example.com/about","https://www.example.com/contact"]}}
    #     - Incorrect: Any additional explanation, markdown, or surrounding text.

    #     OUTPUT FORMAT:
    #     {{"important_links": []}}
    #     """
    # ),
    # ("user", "{url_content}")
    # ])

    prompt = ChatPromptTemplate.from_messages([
    (
        "system",
        """
    You are an expert maritime infrastructure and port-intelligence URL analyzer.

    Your task is to examine the provided Markdown content (generated from crawling a port or port-authority website) and return ONLY the URLs that are highly likely to contain authoritative, factual information relevant to PORT PROFILE, OPERATIONS, GOVERNANCE, INFRASTRUCTURE, FINANCIALS, DIGITAL SYSTEMS, ENVIRONMENT, SAFETY, HR, STRATEGY, or OFFICIAL DOCUMENTS of the port or the organization operating it.

    ========================
    INCLUDE URLs THAT LIKELY CONTAIN INFORMATION RELATED TO:
    ========================

    1) Port Identity & Governance
    - About the Port / About Us / Port Profile
    - Port Authority / Board / Organisation Structure
    - Ministry / Statutory Acts / Major Port Authorities Act
    - History / Establishment
    - Contact details / Official address

    2) Infrastructure Profile
    - Harbour / Channel / Draft / Tidal data
    - Terminals, berths, jetties, cargo handling facilities
    - Equipment, cranes, VTMS, navigation systems
    - Warehousing, storage yards, tank farms, CFS/ICDs

    3) Connectivity & Hinterland
    - Rail, road, inland waterways, pipelines
    - Logistics parks, FTWZ, DFC connectivity

    4) Operational & Performance Data
    - Cargo handled, vessel calls, throughput
    - Productivity, turnaround time, berth occupancy
    - Traffic statistics (annual/monthly)

    5) Digital Systems & Automation
    - PCS 1x, ERP, smart port systems
    - Gati Shakti, automation, RFID, VTMS, IoT

    6) Financial & Commercial Information
    - Tariffs, port dues, pilotage charges
    - Revenue, financial statements, budgets
    - PPP / BOT concession summaries

    7) Environmental & Sustainability
    - Environmental clearance, EIA
    - Sustainability / green port initiatives
    - Pollution monitoring, dredging info

    8) Safety, Compliance & Regulation
    - ISPS, fire safety, disaster management
    - Safety manuals, inspection compliance

    9) Human Resources & Labour
    - Recruitment, HR policies, training institutes
    - Workforce statistics

    10) Stakeholders & Ecosystem
        - Terminal operators, CFS, rail operators
        - Shipping lines, customs, PHO

    11) Strategic Projects & Future Development
        - Sagarmala projects
        - Expansion, modernization plans
        - Master plans, DPRs

    12) Risk, Vulnerability & Cybersecurity
        - Cyber security policies
        - Disaster resilience, climate risk

    13) Publicly Available Official Documents (HIGH PRIORITY)
        - Annual Reports
        - Tariff books
        - Port Master Plans
        - Sustainability Reports
        - SOPs / Manuals
        - Press Releases
        - Government / MoPSW / SPM / NMPA guidelines

    ========================
    EXCLUDE URLs THAT POINT TO:
    ========================
    - News articles unrelated to port operations
    - Tender listings (unless policy / manual)
    - Login pages, dashboards
    - Job application portals (unless HR policy pages)
    - Image galleries, videos
    - Social media links
    - Search, filters, pagination
    - Anchors (#), javascript links
    - Purely transactional services

    ========================
    IMPORTANT RULES (STRICT):
    ========================
    1) Output MUST be valid JSON and NOTHING else.
    2) Output must be exactly:
    {{"important_links": ["https://...", "https://..."]}}
    3) Use ONLY absolute, fully qualified HTTP/HTTPS URLs.
    4) Convert relative links using the base domain of the crawled site.
    5) Deduplicate URLs.
    6) If no qualifying URLs exist, return:
    {{"important_links": []}}

    ========================
    OUTPUT FORMAT:
    ========================
    {{"important_links": []}}
    """
        ),
        ("user", "{url_content}")
    ])
    
    # Chain the prompt, LLM, and parser
    chain = prompt | llm | parser

    # Function to parse product descriptions
    def parse_product(description: str) -> dict:
        result = chain.invoke({"url_content": url_content})
        print("Results : ",result.get("important_links",[]))
        url_links_list = result.get("important_links",[])
        save_docx(url_links_list)
    parse_product(url_content)


if __name__ == "__main__":
    asyncio.run(prompter())